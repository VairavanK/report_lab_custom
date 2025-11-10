import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from .utils import prepare_dataframe, format_value
from .theme import report_colors



# ==================== COLOR CONFIGURATION ====================
# Global color scheme - change once, applies to all sections
report_colors = {
    'header_bg': '#FF8C00',      # Table header background (orange)
    'header_text': '#FFFFFF',    # Table header text (white)
    'row_alt': '#FFE5CC',        # Alternating row background (light orange)
    'row_normal': '#FFFFFF',     # Normal row background (white)
    'separator': '#FF8C00',      # Title separator line (orange)
    'subtitle': '#666666',       # Report subtitle text (grey)
    'section_subtitle': '#888888',  # Section subtitle text (lighter grey)
}

# ==================== DOCX RENDERER FUNCTIONS ====================
def render_docx_title(doc, section_config, theme):
    """Render title section in DOCX"""
    title = section_config.get('title', '')
    subtitle = section_config.get('subtitle', '')

    if title:
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.runs[0].font.size = Pt(14)
        subtitle_para.runs[0].font.italic = True
        subtitle_hex = section_config.get('subtitle_color', theme['subtitle']).replace('#', '')
        subtitle_rgb = tuple(int(subtitle_hex[i:i+2], 16) for i in (0, 2, 4))
        subtitle_para.runs[0].font.color.rgb = RGBColor(*subtitle_rgb)

    doc.add_paragraph()

def render_docx_text(doc, section_config, theme):
    """Render a text paragraph section in DOCX"""
    content = section_config.get('content', '')
    if not content:
        return

    title = section_config.get('title')
    if title:
        doc.add_heading(title, level=2)

    subtitle = section_config.get('subtitle')
    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.runs[0].font.size = Pt(12)
        subtitle_para.runs[0].font.italic = True
        subtitle_hex = theme['section_subtitle'].replace('#', '')
        subtitle_rgb = tuple(int(subtitle_hex[i:i+2], 16) for i in (0, 2, 4))
        subtitle_para.runs[0].font.color.rgb = RGBColor(*subtitle_rgb)

    style = section_config.get('style', 'normal')
    alignment = section_config.get('alignment', 'left')

    paragraph = doc.add_paragraph(content)

    if alignment == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if style == 'bold':
        paragraph.runs[0].font.bold = True
    elif style == 'italic':
        paragraph.runs[0].font.italic = True

    if 'font_size' in section_config:
        paragraph.runs[0].font.size = Pt(section_config['font_size'])

    doc.add_paragraph()

def render_docx_image(doc, section_config, theme):
    """Render an image section in DOCX"""
    image_path = section_config.get('image_path', '')
    if not image_path or not os.path.exists(image_path):
        return

    title = section_config.get('title')
    if title:
        doc.add_heading(title, level=2)

    subtitle = section_config.get('subtitle')
    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.runs[0].font.size = Pt(12)
        subtitle_para.runs[0].font.italic = True
        subtitle_hex = theme['section_subtitle'].replace('#', '')
        subtitle_rgb = tuple(int(subtitle_hex[i:i+2], 16) for i in (0, 2, 4))
        subtitle_para.runs[0].font.color.rgb = RGBColor(*subtitle_rgb)

    caption = section_config.get('caption')
    if caption:
        caption_para = doc.add_paragraph(caption)
        caption_para.runs[0].font.italic = True
        caption_para.runs[0].font.size = Pt(10)

    width = section_config.get('width', 6)
    doc.add_picture(image_path, width=Inches(width))

    doc.add_paragraph()

def render_docx_table(doc, section_config, theme):
    """Render a table section in DOCX"""
    df = section_config['df']
    if len(df) == 0:
        return

    df_work = df.copy()

    if section_config.get('clean_empty_cols'):
        df_work = df_work.dropna(axis=1, how='all')
        df_work = df_work.loc[:, (df_work != '').any(axis=0)]

    if section_config.get('clean_empty_rows'):
        df_work = df_work.dropna(axis=0, how='all')
        df_work = df_work[(df_work.astype(str) != '').any(axis=1)]

    if 'drop_columns' in section_config:
        drop_cols = section_config['drop_columns']
        df_work = df_work.drop(columns=[c for c in drop_cols if c in df_work.columns])

    if len(df_work) == 0:
        return

    df_work = prepare_dataframe(df_work)

    title = section_config.get('title', '')
    if 'title_suffix_from_column' in section_config:
        col = section_config['title_suffix_from_column']
        if col in df.columns:
            val = df[col].iloc[0]
            title = f"{title} ({val})"

    if title:
        doc.add_heading(title, level=2)

    subtitle = section_config.get('subtitle')
    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.runs[0].font.size = Pt(12)
        subtitle_para.runs[0].font.italic = True
        subtitle_hex = theme['section_subtitle'].replace('#', '')
        subtitle_rgb = tuple(int(subtitle_hex[i:i+2], 16) for i in (0, 2, 4))
        subtitle_para.runs[0].font.color.rgb = RGBColor(*subtitle_rgb)

    section_colors = section_config.get('colors', theme)
    header_hex = section_colors['header_bg'].replace('#', '')
    header_rgb = tuple(int(header_hex[i:i+2], 16) for i in (0, 2, 4))
    row_alt_hex = section_colors['row_alt'].replace('#', '')
    row_alt_rgb = tuple(int(row_alt_hex[i:i+2], 16) for i in (0, 2, 4))

    table_data = [df_work.columns.tolist()]
    for _, row in df_work.iterrows():
        table_data.append([format_value(v) for v in row])

    table = doc.add_table(rows=len(table_data), cols=len(df_work.columns))
    table.style = 'Table Grid'

    for j, cell in enumerate(table.rows[0].cells):
        cell.text = str(table_data[0][j])
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), section_colors['header_bg'].replace('#', ''))
        cell._element.get_or_add_tcPr().append(shading_elm)

    for i, row_data in enumerate(table_data[1:], start=1):
        for j, cell in enumerate(table.rows[i].cells):
            cell.text = str(row_data[j])
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            if i % 2 == 0:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), section_colors['row_alt'].replace('#', ''))
                cell._element.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()

def render_docx_table_grouped(doc, section_config, theme):
    """Render grouped tables in DOCX"""
    df = section_config['df']
    if len(df) == 0:
        return

    groupby_col = section_config['groupby']
    if groupby_col not in df.columns:
        return

    subtitle = section_config.get('subtitle')
    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.runs[0].font.size = Pt(12)
        subtitle_para.runs[0].font.italic = True
        subtitle_hex = theme['section_subtitle'].replace('#', '')
        subtitle_rgb = tuple(int(subtitle_hex[i:i+2], 16) for i in (0, 2, 4))
        subtitle_para.runs[0].font.color.rgb = RGBColor(*subtitle_rgb)

    category_order = section_config.get('category_order', df[groupby_col].unique())

    for category in category_order:
        df_cat = df[df[groupby_col] == category]
        if len(df_cat) == 0:
            continue

        section_copy = section_config.copy()
        section_copy['title'] = category
        section_copy['df'] = df_cat
        section_copy['type'] = 'table'

        render_docx_table(doc, section_copy, theme)

def _generate_docx_report(report_sections, output_filename, theme):
    """Generate DOCX report"""
    doc = Document()

    for section in report_sections:
        section_type = section['type']

        if section_type == 'title':
            render_docx_title(doc, section, theme)
        elif section_type == 'text':
            render_docx_text(doc, section, theme)
        elif section_type == 'image':
            render_docx_image(doc, section, theme)
        elif section_type == 'table':
            render_docx_table(doc, section, theme)
        elif section_type == 'table_grouped':
            render_docx_table_grouped(doc, section, theme)

    doc.save(output_filename)
    return output_filename